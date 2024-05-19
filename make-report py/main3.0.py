import copy
import threading
import time
import sys
import os
import itertools
import excel2img
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup as bs
import datetime
import requests
from openpyxl import Workbook, load_workbook
from openpyxl.styles import *
from docx import Document
from docx.shared import Mm

def loading_animation(stop_event, message):
    global load_ok
    chars = itertools.cycle(["|", "/", "-", "\\"])
    while not stop_event.is_set():
        sys.stdout.write("\r" + f'{message} ' + next(chars))
        sys.stdout.flush()
        time.sleep(0.1)
    if load_ok:
        sys.stdout.write("\r" + f'{message} ✔\n')
    else:
        sys.stdout.write("\r" + f'{message} х\n')
    sys.stdout.flush()
    load_ok = True

def is_number(str):
    try:
        float(str)
        return True
    except ValueError:
        return False

def get_data(d,m,y):
    H1 = 9
    H2 = 17
    url_in = f'http://dbrobo.mgul.ac.ru/core/deb.php?fdate={y}-{m}-{d}+{H1}%3A00%3A00&sdate={y}-{m}-{d}+{H2}%3A00%3A00&fileback=1'
    url_out = "http://dbrobo.mgul.ac.ru/export/log.txt"
    requests.get(url_in)
    r = requests.get(url_out)
    soup = bs(r.text, "lxml")
    data = eval(soup.text)
    return data


def get_devices(data):
    devices = []
    dictionary = {}
    pattern = {}
    pattern_res = {}
    pattern_stat = {}
    cur_time = datetime.time(9, 0, 0)
    for col in range(16):
        pattern[cur_time.strftime("%H:%M")] = []
        pattern_res[cur_time.strftime("%H:%M")] = None
        pattern_stat[cur_time.strftime("%H:%M")] = 0
        try:
            cur_time = cur_time.replace(minute=cur_time.minute + 30)
        except:
            cur_time = cur_time.replace(hour=cur_time.hour + 1, minute=0)

    for el in data:
        devices.append(data[el]['uName'] + ' (' + data[el]['serial'] + ')')
    devices = sorted(list(set(devices)))
    i = 0
    while i < len(devices):
        if 'Сервер' in devices[i]:
            del devices[i]
        else:
            i += 1

    for device in devices:
        dictionary[device] = {}
        dictionary[device]['result'] = {}
        if 'Hydra-L' in device:
            dictionary[device]['system_RSSI'] = copy.deepcopy(pattern)
            dictionary[device]['BME280_temp'] = copy.deepcopy(pattern)
            dictionary[device]['BME280_humidity'] = copy.deepcopy(pattern)
            dictionary[device]['BME280_pressure'] = copy.deepcopy(pattern)

        elif 'Опорный барометр' in device:
            dictionary[device]['system_RSSI'] = copy.deepcopy(pattern)
            for i in range(9):
                dictionary[device][f'BMP280_{i}_temp'] = copy.deepcopy(pattern)
                dictionary[device][f'BMP280_{i}_pressure'] = copy.deepcopy(pattern)

            dictionary[device]['weather_temp'] = copy.deepcopy(pattern)
            dictionary[device]['weather_temp_mediana'] = copy.deepcopy(pattern)
            dictionary[device]['weather_pressure'] = copy.deepcopy(pattern)
            dictionary[device]['weather_pressure_mediana'] = copy.deepcopy(pattern)

        elif 'Паскаль' in device:
            dictionary[device]['system_RSSI'] = copy.deepcopy(pattern)
            dictionary[device]['weather_temp'] = copy.deepcopy(pattern)
            dictionary[device]['weather_pressure'] = copy.deepcopy(pattern)

        elif 'РОСА К-2' in device or 'Роса-К-1' in device:
            dictionary[device]['system_RSSI'] = copy.deepcopy(pattern)
            dictionary[device]['color_tempCT'] = copy.deepcopy(pattern)
            #dictionary[device]['soil_soilT'] = copy.deepcopy(pattern)
            dictionary[device]['weather_temp'] = copy.deepcopy(pattern)
            dictionary[device]['weather_pressure'] = copy.deepcopy(pattern)
            dictionary[device]['weather_humidity'] = copy.deepcopy(pattern)
            dictionary[device]['light_lux'] = copy.deepcopy(pattern)
            dictionary[device]['light_blink'] = copy.deepcopy(pattern)

        elif 'Тест Студии' in device:
            dictionary[device]['system_RSSI'] = copy.deepcopy(pattern)
            dictionary[device]['TCS34725_luxCCT'] = copy.deepcopy(pattern)
            dictionary[device]['TCS34725_colorTempCT'] = copy.deepcopy(pattern)
            dictionary[device]['SBM20_static'] = copy.deepcopy(pattern)
            dictionary[device]['SBM20_dynamic'] = copy.deepcopy(pattern)
            dictionary[device]['DS18B20_temp'] = copy.deepcopy(pattern)
            dictionary[device]['CCS811_eCO2'] = copy.deepcopy(pattern)
            dictionary[device]['CCS811_TVOC'] = copy.deepcopy(pattern)
            dictionary[device]['CCS811_ErrFlag'] = copy.deepcopy(pattern)
            dictionary[device]['CCS811_ErrCode'] = copy.deepcopy(pattern)
            dictionary[device]['BMP280_temp'] = copy.deepcopy(pattern)
            dictionary[device]['BMP280_pressure'] = copy.deepcopy(pattern)
            dictionary[device]['BME280_temp'] = copy.deepcopy(pattern)
            dictionary[device]['BME280_pressure'] = copy.deepcopy(pattern)
            dictionary[device]['BME280_humidity'] = copy.deepcopy(pattern)
            dictionary[device]['BH1750_lux'] = copy.deepcopy(pattern)
            dictionary[device]['BH1750_blink'] = copy.deepcopy(pattern)
            dictionary[device]['AM2321_temp'] = copy.deepcopy(pattern)
            dictionary[device]['AM2321_humidity'] = copy.deepcopy(pattern)

        elif 'Тест воздуха' in device:
            dictionary[device]['system_RSSI'] = copy.deepcopy(pattern)
            dictionary[device]['MHZ19B_CO2'] = copy.deepcopy(pattern)
            dictionary[device]['CCS811_eCO2'] = copy.deepcopy(pattern)
            dictionary[device]['CCS811_TVOC'] = copy.deepcopy(pattern)
            dictionary[device]['BME280_temp'] = copy.deepcopy(pattern)
            dictionary[device]['BME280_humidity'] = copy.deepcopy(pattern)
            dictionary[device]['BME280_pressure'] = copy.deepcopy(pattern)

        dictionary[device]['result']['status'] = copy.deepcopy(pattern_stat)

        for sensor in dictionary[device]:
            if sensor != 'result':
                dictionary[device]['result'][sensor] = copy.deepcopy(pattern_res)
                dictionary[device]['result'][sensor + '_error'] = copy.deepcopy(pattern_res)


    return dictionary

def processing(data, base):
    cur_time = datetime.time(9, 0, 0)
    next_time = datetime.time(9, 30, 0)
    cur_time_str = cur_time.strftime("%H:%M")
    next_time_str = next_time.strftime("%H:%M")

    # наполнение base из файла
    for elem in data:
        device = data[elem]['uName'] + ' (' + data[elem]['serial'] + ')'
        if device in base:
            try:
                d_time = datetime.time.fromisoformat(data[elem]['Date'].split(' ')[1])
            except:
                # print("Ошибка в дате", data[elem]['Date'])
                continue
            if d_time > next_time:
                try:
                    cur_time = next_time
                    next_time = next_time.replace(minute=next_time.minute + 30)
                except:
                    cur_time = next_time
                    next_time = next_time.replace(hour=next_time.hour + 1, minute=0)
                cur_time_str = cur_time.strftime("%H:%M")
                next_time_str = next_time.strftime("%H:%M")
            for sensor in data[elem]['data']:
                if sensor in base[device]:
                    print(data[elem]['data'][sensor], sensor)
                    cur = float(data[elem]['data'][sensor])
                    base[device][sensor][cur_time_str].append(cur)

    #тут считается кол-во записей за пол часа и выдаётся значение в статус
    for device in base:
        for sensor in base[device]:
            if sensor != 'result':
                for segment in base[device][sensor]:
                    records = len(base[device][sensor][segment])
                    # статус = -1 если нет сигнала за пол часа
                    if records == 0:
                        base[device]['result']['status'][segment] = -1
                        continue
                    # если записей меньше нормы, статус = процент полученных записей от нормы
                    if 'Hydra-L' in device:
                        print(records)
                        if records < 30:
                            base[device]['result']['status'][segment] = int(records * 100 / 30 + 0.5)

                    elif 'Опорный барометр' in device:
                        if '1' in device != -1:
                            if records < 29:
                                base[device]['result']['status'][segment] = int(records * 100 / 29 + 0.5)
                        if '2' in device != -1:
                            if records < 28:
                                base[device]['result']['status'][segment] = int(records * 100 / 28 + 0.5)

                    elif 'Паскаль' in device:
                        if records < 27:
                            base[device]['result']['status'][segment] = int(records * 100 / 27 + 0.5)

                    elif 'РОСА К-2' in device or 'Роса-К-1' in device:
                        if records < 28:
                            base[device]['result']['status'][segment] = int(records * 100 / 28 + 0.5)

                    elif 'Тест Студии' in device:
                        if records < 29:
                            base[device]['result']['status'][segment] = int(records * 100 / 29 + 0.5)

                    elif 'Тест воздуха' in device:
                        if records < 30:
                            base[device]['result']['status'][segment] = int(records * 100 / 30 + 0.5)

                break


    # идея в том чтобы отдельно пойтись по массивам и добавить крайние элементы вручную
    flag_first = False
    last_arr = []
    last_elem = 0
    # print(base)
    for device in base:
        for sensor in base[device]:
            if sensor != 'result':
                for segment in base[device][sensor]:
                    if segment == "09:00":
                        if len(base[device][sensor][segment]) > 0:
                            base[device][sensor][segment].insert(0, base[device][sensor][segment][0])
                        else:
                            flag_first = True
                    if len(base[device][sensor][segment]) > 0:
                        if flag_first:
                            base[device][sensor][segment].insert(0, base[device][sensor][segment][0])
                            flag_first = False
                        else:
                            if segment != "09:00":
                            # добавить последним первый элемент из следующего полного массива
                                last_arr.append(base[device][sensor][segment][0])
                            # добавить первым элемент из последнего полного массива
                                base[device][sensor][segment].insert(0, first_elem)
                        last_arr = base[device][sensor][segment]
                        first_elem = last_arr[-1]
                last_arr.append(first_elem)




    #генерация значений выдачи в эксель
    for device in base:
        for sensor in base[device]:
            if sensor != 'result':
                for segment in base[device][sensor]:
                    val_arr = base[device][sensor][segment] #массив значений за пол часа
                    #base[device]['result'][sensor][segment] значение выдачи в эксель
                    if len(val_arr[1:-1]) == 0: #нет показаний
                        base[device]['result'][sensor][segment] = 'F'
                    elif len(val_arr[1:-1]) != 1 and len(set(val_arr[1:-1])) == 1: #неизменяемость значений
                        if device == 'Тест Студии (schHome)' and (sensor == 'CCS811_ErrFlag' or sensor == 'BH1750_blink' or sensor == 'TCS34725_luxCCT')\
                                or sensor == 'system_RSSI' or sensor == 'CCS811_ErrCode':
                            base[device]['result'][sensor][segment] = round(sum(val_arr[1:-1]) / len(val_arr[1:-1]), 2)
                        else:
                            base[device]['result'][sensor][segment] = val_arr[1]
                            base[device]['result'][sensor + '_error'][segment] = 'E'
                    else:
                        #поиск маргиналов
                        flag = 0
                        for i in range(1, len(val_arr)-1):
                            try:
                                prev, cur, foll = val_arr[i-1], val_arr[i], val_arr[i+1]
                                c1 = (abs((cur - foll)) / abs(foll)) * 100
                                c2 = (abs((foll - cur)) / abs(cur)) * 100
                                c3 = (abs((prev - cur)) / abs(cur)) * 100
                                c4 = (abs((cur - prev)) / abs(prev)) * 100
                                check, check2 = max(c1, c2), max(c3, c4)

                                # если 1 чек >30 то резкая смена тренда, если 2 то единичный элемент
                                if check > 30 and check2 > 30:
                                    # print(f'Маргинал найден {device} {sensor} {prev} {cur} {foll}')
                                    #CCS811_ErrFlag значения 0/1, BH1750_blink от 0 до 100, TCS34725_luxCCT и CCS811_TVOC колбасит
                                    if device == 'Тест Студии (schHome)' and (
                                            sensor == 'CCS811_ErrFlag' or sensor == 'BH1750_blink' or sensor == 'TCS34725_luxCCT' or sensor == 'CCS811_TVOC'):
                                        pass
                                    elif device == 'Тест Студии (schHome)' and sensor == 'BH1750_lux':
                                        if check > 95 and check2 > 95:
                                            flag += 1
                                            val_arr.pop(i)
                                    elif 'Тест воздуха' in device and sensor == 'CCS811_TVOC':
                                        if check > 100 and check2 > 100:
                                            flag += 1
                                            val_arr.pop(i)
                                    elif sensor == 'system_RSSI':
                                        pass
                                    elif abs(prev) < 3 and abs(cur) < 3: #(опасность при смене знаков)
                                        pass
                                    else:
                                        flag += 1
                                        val_arr.pop(i)
                                if 'pressure' in sensor and (check > 5 and check2 > 5):
                                    flag += 1
                                    val_arr.pop(i)
                            except:
                                pass
                                # print(f'Нулевое значение {device} {sensor} {segment} пропущено')

                        if flag > 0:
                            base[device]['result'][sensor + '_error'][segment] = f'A/B ({flag})'

                        base[device]['result'][sensor][segment] = round(sum(val_arr[1:-1]) / len(val_arr[1:-1]), 2)

    copy_base = copy.deepcopy(base)
    #если ошибок нет удаляем словарь из выдачи
    for device in copy_base:
        for sensor in copy_base[device]['result']:
            if '_error' in sensor:
                flag = True
                for segment in copy_base[device]['result'][sensor]:
                    if copy_base[device]['result'][sensor][segment] != None:
                        flag = False
                if flag:
                    del base[device]['result'][sensor]



#A	Невозможное значение
#B	Единичное значение, выпадающее из последовательности
#C	Разрыв
#D	Слишком частое и резкое изменение значений и их производных
#E	Неизменяемость значений
#F	Отсутствие показаний
#'Тест Студии (schHome)', 'CCS811_TVOC' скочит сильно
#BH1750_blink скочит от 0 до 100 моментами
#'РОСА К-2 (01)' 'soil_soilT' от - к + не больше |1|
#Тест Студии (schHome)', 'DS18B20_temp', от - к + не больше |2|



def excel_out(base, filename):
    device_list_pc = ['Hydra-L (01)', 'Hydra-L (02)', 'Hydra-L (03)', 'Hydra-L (04)', 'Hydra-L (05)', 'Hydra-L (06)', 'Hydra-L (08)', 'Hydra-L1 (01)', 'Hydra-L1 (02)', 'РОСА К-2 (01)', 'Опорный барометр (01)']
    device_list_class = ['Hydra-L (07)']

    wb = Workbook()
    ws = wb.active
    thins = Side(border_style="thin", color="000000")

    my_green = colors.Color(rgb='0000FF00')
    my_fill_green = fills.PatternFill(patternType='solid', fgColor=my_green)
    my_yellow = colors.Color(rgb='00FFFF00')
    my_fill_yellow = fills.PatternFill(patternType='solid', fgColor=my_yellow)
    my_red = colors.Color(rgb='00FF0000')
    my_fill_red = fills.PatternFill(patternType='solid', fgColor=my_red)

    my_low = colors.Color(rgb='00CCECFF')
    my_fill_low = fills.PatternFill(patternType='solid', fgColor=my_low)
    my_high = colors.Color(rgb='00FFFFCC')
    my_fill_high = fills.PatternFill(patternType='solid', fgColor=my_high)

    my_ex = colors.Color(rgb='0066FF33')
    my_fill_ex = fills.PatternFill(patternType='solid', fgColor=my_ex)
    my_good = colors.Color(rgb='00CCFF33')
    my_fill_good = fills.PatternFill(patternType='solid', fgColor=my_good)
    my_fair = colors.Color(rgb='00FFCC00')
    my_fill_fair = fills.PatternFill(patternType='solid', fgColor=my_fair)
    my_poor = colors.Color(rgb='00FF6600')
    my_fill_poor = fills.PatternFill(patternType='solid', fgColor=my_poor)
    my_no = colors.Color(rgb='00FF0000')
    my_fill_no = fills.PatternFill(patternType='solid', fgColor=my_no)

    flag = True
    for device in base:
        if flag:
            flag = False
            ws.title = device
        else:
            ws = wb.create_sheet(device)

        ws.cell(row=1, column=1).border = Border(bottom=thins, right=thins)

        cur_time = datetime.time(9, 0, 0)

        for col in range(2, 18):
            cell_time = ws.cell(row=1, column=col)
            cell_time.value = cur_time.strftime("%H:%M")
            cell_time.alignment = Alignment(horizontal='center')
            cell_time.font = Font(bold=True)
            if col == 17:
                cell_time.border = Border(bottom=thins, right=thins)
            else:
                cell_time.border = Border(bottom=thins)
            try:
                cur_time = cur_time.replace(minute=cur_time.minute + 30)
            except:
                cur_time = cur_time.replace(hour=cur_time.hour + 1, minute=0)

        r = 2
        ws.column_dimensions['A'].width = 25
        for sensor in base[device]['result']:
            name_cell = ws.cell(row=r, column=1)
            if sensor == 'status':
                name_cell.value = 'Статус'
            else:
                name_cell.value = sensor
            name_cell.alignment = Alignment(horizontal='left')
            if r == len(base[device]['result']) + 1:
                name_cell.border = Border(right=thins, bottom=thins)
            else:
                name_cell.border = Border(right=thins)
            i = 2
            for segment in base[device]['result'][sensor]:
                cell = ws.cell(row=r, column=i)
                v = base[device]['result'][sensor][segment]
                if sensor == 'status':
                    if v == 0:
                        cell.fill = my_fill_green
                    elif v == -1:
                        cell.fill = my_fill_red
                    else:
                        cell.fill = my_fill_yellow
                        cell.value = v / 100
                        cell.number_format = numbers.BUILTIN_FORMATS[9]
                else:
                    cell.value = v
                    if v != None and not is_number(v):
                        cell.font = Font(color="FF0000", bold=True)
                        cell.alignment = Alignment(horizontal='center')
                    else:
                        if 'RSSI' in sensor:
                            if v > -70:
                                cell.fill = my_fill_ex
                            elif -70 >= v >= -85:
                                cell.fill = my_fill_good
                            elif -85 > v >= -100:
                                cell.fill = my_fill_fair
                            elif -100 > v > -110:
                                cell.fill = my_fill_poor
                            else:
                                cell.fill = my_fill_no
                        if device in device_list_pc + device_list_class:
                            if 'error' not in sensor:
                                if 'pressure' in sensor:
                                    if v < 740:
                                        cell.fill = my_fill_low
                                    if v > 750:
                                        cell.fill = my_fill_high
                                elif 'temp' in sensor and 'color' not in sensor:
                                    if v < 18:
                                        cell.fill = my_fill_low
                                    if v > 24:
                                        cell.fill = my_fill_high
                                elif 'humidity' in sensor:
                                    if device in device_list_class:
                                        if v < 40:
                                            cell.fill = my_fill_low
                                        elif v > 60:
                                            cell.fill = my_fill_high
                                    else:
                                        if v < 55:
                                            cell.fill = my_fill_low
                                        elif v > 62:
                                            cell.fill = my_fill_high
                if r == len(base[device]['result']) + 1:
                    cell.border = Border(bottom=thins)
                    if i == 17:
                        cell.border = Border(bottom=thins, right=thins)
                    else:
                        cell.border = Border(bottom=thins)
                else:
                    if i == 17:
                        cell.border = Border(right=thins)
                i += 1
            r += 1

    if not os.path.exists('Отчёты'):
        os.mkdir("Отчёты")
    wb.save(f'Отчёты/{filename}.xlsx')


def excel_to_word(filename):
    devices = {
        'Hydra-L (01)': 0,
        'Hydra-L (02)': 0,
        'Hydra-L (03)': 0,
        'Hydra-L (04)': 0,
        'Hydra-L (05)': 0,
        'Hydra-L (06)': 0,
        'Hydra-L (07)': 0,
        'Hydra-L (08)': 0,
        'Hydra-L1 (01)': 0,
        'Hydra-L1 (02)': 0,
        'Опорный барометр (01)': 0,
        'Опорный барометр (02)': 0,
        'Паскаль (00)': 0,
        'Паскаль (01)': 0,
        'Паскаль (02)': 0,
        'Паскаль (03)': 0,
        'Паскаль (04)': 0,
        'Паскаль (05)': 0,
        'Паскаль (06)': 0,
        'Паскаль (07)': 0,
        'Паскаль (08)': 0,
        'Паскаль (09)': 0,
        'Паскаль (10)': 0,
        'Паскаль (11)': 0,
        'Паскаль (12)': 0,
        'Паскаль (13)': 0,
        'Паскаль (14)': 0,
        'Паскаль (15)': 0,
        'Паскаль (16)': 0,
        'Паскаль (17)': 0,
        'Паскаль (18)': 0,
        'Паскаль (19)': 0,
        'Паскаль (20)': 0,
        'РОСА К-2 (01)': 0,
        'Роса-К-1 (01)': 0,
        'Тест Студии (schHome)': 0,
        'Тест воздуха (01)': 0
    }

    doc = Document('Отчёты/шаблон отчёта.docx')

    styles = doc.styles

    try:
        style = styles.add_style('style_text', WD_STYLE_TYPE.CHARACTER)
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
    except:
        pass
    try:
        style = styles.add_style('style_device', WD_STYLE_TYPE.CHARACTER)
        font = style.font
        font.name = 'Cambria'
        font.italic = True
        font.size = Pt(14)
    except:
        pass
    try:
        style = styles.add_style('style_paragraph_device', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Cambria'
        font.italic = True
        font.size = Pt(14)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(10)
        style.paragraph_format.keep_with_next = True
    except:
        pass

    workbook = load_workbook(f'Отчёты/{filename}.xlsx')
    for sheet in workbook.sheetnames:
        devices[sheet] = 1

    p = doc.add_paragraph()
    p.add_run('Приборы ', style='style_text')
    p.add_run((lambda d: ', '.join(key for key, value in d.items() if value == 0))(devices), style='style_device')
    p.add_run(' не выходили на связь на протяжении всей смены.', style='style_text')
    doc.add_paragraph().add_run('Приборы ... выдавали значительные/незначительные аномалии.', style='style_text')
    p = doc.add_paragraph()
    p.add_run('Прибор ', style='style_text')
    p.add_run('Тест Студии (schHome) ', style='style_device')
    p.add_run('работал некорректно.', style='style_text')

    doc.add_page_break()

    for sheet in workbook.sheetnames:
        excel2img.export_img(f'Отчёты/{filename}.xlsx', 'Отчёты/image.png', sheet)
        doc.add_paragraph(f'{sheet}', style='style_paragraph_device')
        doc.add_picture('Отчёты/image.png', width=Mm(165))

    os.remove('Отчёты/image.png')
    doc.save(f'Отчёты/{filename} отчёт 2 бригады по 2 заданию.docx')


def yes_no_dialog(question, r=0):
    answers = {"yes":1, "y":1, "ye":1, "н":1,
        "no":0, "n":0, "т":0}
    tip = " [y/n] "
    while True:
        print(question + tip + ": ", end="")
        user_answer = input().lower()
        if user_answer in answers:
            if answers[user_answer] == 1 and r == 0:
                if datetime.datetime.now().hour < 17:
                    print("Ещё рано формировать отчёт за сегодня, дождитесь 17:00")
                else:
                    return answers[user_answer]
            else:
                return answers[user_answer]
        else:
            print("Пожалуйста, введите yes/y или no/n\n")

def user_date(question):
    print(question)
    while True:
        print("День ", end="")
        d = int(input())
        print("Месяц ", end="")
        m = int(input())
        print("Год ", end="")
        y = int(input())
        try:
            if datetime.datetime.now() > datetime.datetime(y, m, d, 17, 0, 0):
                return d, m, y
            else:
                print("Для данной даты ещё рано формировать отчёт, выберите другую или подождите")
        except:
            print("Ошибка в дате, повторите попытку")

load_ok = True
def main():
    global load_ok
    if yes_no_dialog("Сгенерировать отчёт за сегодня?"):
        d = datetime.date.today().day
        m = datetime.date.today().month
        y = datetime.date.today().year
    else:
        d, m, y = user_date("Введите дату для формирования отчёта")

    date_string = datetime.datetime(y, m, d).strftime("%d.%m.%Y")

    data, base = {}, {}
    while True:
        stop_event = threading.Event()
        loading_thread = threading.Thread(target=loading_animation, args=(stop_event, 'Загрузка данных с сервера'))
        loading_thread.start()
        try:
            data = get_data(d,m,y)
            stop_event.set()
            loading_thread.join()
            break
        except:
            load_ok = False
            stop_event.set()
            loading_thread.join()
            if not yes_no_dialog("Данные с сервера не получены, попробуйте повторить попытку?", 1):
                return

    stop_event = threading.Event()
    loading_thread = threading.Thread(target=loading_animation, args=(stop_event, 'Обработка информации'))
    loading_thread.start()

    base = get_devices(data)
    processing(data, base)

    stop_event.set()
    loading_thread.join()
    excel_out(base, date_string)
    while True:
        try:
            stop_event = threading.Event()
            loading_thread = threading.Thread(target=loading_animation, args=(stop_event, 'Создание excel файла'))
            loading_thread.start()

            excel_out(base, date_string)

            stop_event.set()
            loading_thread.join()
            break
        except Exception as e:
            print(e)
            load_ok = False
            stop_event.set()
            loading_thread.join()
            if not yes_no_dialog("Закройте excel файл с отчётом! Повторить попытку?", 1):
                return

    while True:
        try:
            stop_event = threading.Event()
            loading_thread = threading.Thread(target=loading_animation, args=(stop_event, 'Создание word файла'))
            loading_thread.start()

            excel_to_word(date_string)

            stop_event.set()
            loading_thread.join()
            break
        except Exception as e:
            print(e)
            load_ok = False
            stop_event.set()
            loading_thread.join()
            if not yes_no_dialog("Закройте word файл с отчётом! Повторить попытку?", 1):
                return


    print(f'Отчёт за {date_string} сгенерирован')

if __name__ == '__main__':
    main()
    input()


